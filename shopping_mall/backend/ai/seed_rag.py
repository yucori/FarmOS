"""ChromaDB RAG 데이터 적재 스크립트.

실행: uv run python ai/seed_rag.py

정책 문서(PDF/DOCX)는 DOCS_DIR 경로에 위치해야 합니다.
새 문서 추가 시 DOC_TO_COLLECTION에 파일명과 컬렉션명 매핑을 추가하세요.
"""
import json
import os
import re
import sys

# shopping_mall/backend를 sys.path에 추가 (python ai/seed_rag.py 직접 실행 시 필요)
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from app.core.config import settings
from app.paths import CHROMA_DB_PATH, AI_DATA_DIR
from ai.embeddings import get_embedding_function

DATA_DIR = str(AI_DATA_DIR)
CHROMA_DIR = CHROMA_DB_PATH
DOCS_DIR = settings.policy_docs_dir

# BM25 인덱스 저장 경로 (ai/data/bm25_index.json)
BM25_INDEX_PATH = str(AI_DATA_DIR / "bm25_index.json")

# 시딩 대상 전체 컬렉션 — BM25 인덱스 빌드 시 사용
_ALL_COLLECTIONS = [
    "faq",
    "storage_guide",
    "season_info",
    "farm_intro",
    "payment_policy",
    "delivery_policy",
    "return_policy",
    "quality_policy",
    "service_policy",
    "membership_policy",
]


def _tokenize_ko(text: str) -> list[str]:
    """한국어 정규식 토크나이저 — 한글/영문/숫자 단위로 분리."""
    tokens = re.findall(r"[가-힣a-zA-Z0-9]+", text.lower())
    return tokens if tokens else [text.lower()]


def build_bm25_index(client, collection_names: list[str]) -> dict:
    """ChromaDB 컬렉션에서 전체 문서를 읽어 BM25 인덱스 데이터를 구성.

    Returns:
        {
            "corpus": [[token, ...], ...],  # 토큰화된 문서 목록
            "ids":    ["doc_id", ...],       # 문서 ID (corpus와 1:1 대응)
            "collections": ["col", ...],     # 각 문서의 소속 컬렉션
        }
    """
    corpus: list[list[str]] = []
    ids: list[str] = []
    cols: list[str] = []

    for col_name in collection_names:
        try:
            col = client.get_collection(name=col_name)
            res = col.get(include=["documents"])
            for doc, doc_id in zip(res.get("documents", []), res.get("ids", [])):
                corpus.append(_tokenize_ko(doc))
                ids.append(doc_id)
                cols.append(col_name)
        except Exception as e:
            print(f"  [BM25] {col_name} 읽기 실패: {e}")

    return {"corpus": corpus, "ids": ids, "collections": cols}

# 파일명(부분 일치) → ChromaDB 컬렉션명 매핑
DOC_TO_COLLECTION: dict[str, str] = {
    "01_주문및결제정책": "payment_policy",
    "02_배송정책": "delivery_policy",
    "03_반품교환환불정책": "return_policy",
    "04_상품품질신선도보증정책": "quality_policy",
    "05_고객서비스운영정책": "service_policy",
    "06_개인정보처리및회원정책": "membership_policy",
}


# ── 문서 파싱 ──────────────────────────────────────────────────────────────

def _table_data_to_md(data: list[list]) -> str:
    """2D 리스트 → 마크다운 파이프 테이블 (첫 행을 헤더로 처리)."""
    rows = [
        [str(c or "").strip().replace("\n", " ") for c in row]
        for row in data if any(c for c in row)
    ]
    if not rows:
        return ""
    header = rows[0]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _rect_overlap(a: tuple, b: tuple) -> bool:
    """두 (x0,y0,x1,y1) 사각형 겹침 판정."""
    return not (a[2] <= b[0] or a[0] >= b[2] or a[3] <= b[1] or a[1] >= b[3])


def _apply_heading_markdown(text: str) -> str:
    """제N장/제N조 패턴 줄에 마크다운 헤딩 마커를 적용한다.

    PDF/DOCX 공통 후처리:
      제N장 ...          → ## 제N장 ...   (bare)
      제N조(...)         → ### 제N조(...) (bare — PDF blocks 추출 결과)
      **제N조(...)**     → ### 제N조(...) (DOCX Bold 스타일 결과)
    줄 단위 연속 공백도 이 단계에서 정리한다.
    """
    lines = [re.sub(r"[ \t]{2,}", " ", line.strip()) for line in text.split("\n")]
    text = "\n".join(lines)
    # 제N장: bare 또는 기존 # 헤딩(DOCX Heading 스타일) 모두 ## 으로 통일
    text = re.sub(r"^#{0,6}\s*(제\d+장(?:\s+.+)?)$", r"## \1", text, flags=re.MULTILINE)
    # 제N조: bare(PDF) 또는 **-wrapped(DOCX Bold) 모두 ### 으로 통일
    text = re.sub(r"^\**(제\d+조\(.+?\).*?)\**$", r"### \1", text, flags=re.MULTILINE)
    return text


def parse_pdf(path: str) -> str:
    """PDF → 마크다운 (pymupdf 블록 추출 + 표 감지).

    - find_tables(): 표를 감지해 파이프 테이블로 변환
    - get_text("blocks"): 본문 텍스트 블록을 Y 좌표 순으로 추출 (표 영역 제외)
    - _apply_heading_markdown(): 제N장/제N조 패턴에 ## / ### 헤딩 마커 부여
    """
    import fitz
    doc = fitz.open(path)
    page_parts: list[str] = []

    for page in doc:
        tabs = page.find_tables()
        table_rects = [t.bbox for t in tabs.tables]
        table_items: list[tuple[float, str]] = []
        for t in tabs.tables:
            data = t.extract()
            md = _table_data_to_md(data) if data else ""
            if md:
                table_items.append((t.bbox[1], md))

        text_items: list[tuple[float, str]] = []
        for block in page.get_text("blocks"):
            bx0, by0, bx1, by1, text = block[:5]
            block_type = block[6]
            if block_type != 0:  # 0=text, 1=image
                continue
            if any(_rect_overlap((bx0, by0, bx1, by1), tr) for tr in table_rects):
                continue
            text = text.strip()
            if text:
                text_items.append((by0, text))

        all_items = sorted(text_items + table_items, key=lambda x: x[0])
        page_parts.append("\n\n".join(t for _, t in all_items))

    doc.close()
    return "\n\n".join(page_parts)


# DOCX 헤딩 스타일명 패턴 (python-docx는 "Heading 1", "heading 2" 등으로 반환)
_DOCX_HEADING_RE = re.compile(r"^heading\s*(\d+)$", re.IGNORECASE)


def _runs_to_md(para) -> str:
    """단락 runs → 마크다운 인라인 (볼드 run은 **text** 로 변환)."""
    parts = []
    for run in para.runs:
        if not run.text:
            continue
        parts.append(f"**{run.text}**" if run.bold else run.text)
    return "".join(parts).strip()


def parse_docx(path: str) -> str:
    """DOCX → 마크다운 (python-docx, 단락·표 문서 순서 유지).

    - Heading N 스타일 → # (N개)
    - 볼드 run → **text**
    - 표 → | cell | ... | 마크다운 표
    """
    from docx import Document
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = Document(path)
    parts: list[str] = []

    for child in doc.element.body:
        tag = child.tag
        if tag == qn("w:p"):
            para = Paragraph(child, doc)
            style_name = para.style.name if para.style else ""
            m = _DOCX_HEADING_RE.match(style_name)
            if m:
                level = min(int(m.group(1)), 6)
                text = para.text.strip()
                if text:
                    parts.append(f"{'#' * level} {text}")
            else:
                line = _runs_to_md(para)
                if line:
                    parts.append(line)
        elif tag == qn("w:tbl"):
            tbl = Table(child, doc)
            data = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in tbl.rows]
            md_tbl = _table_data_to_md(data)
            if md_tbl:
                parts.append(md_tbl)

    return "\n\n".join(parts)


def parse_document(path: str) -> str:
    """확장자에 따라 PDF 또는 DOCX 파싱 후 마크다운 헤딩 후처리."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        raw = parse_pdf(path)
    elif ext == ".docx":
        raw = parse_docx(path)
    else:
        raise ValueError(f"지원하지 않는 파일 형식: {ext}")
    return _apply_heading_markdown(raw)


# ── 청킹 ──────────────────────────────────────────────────────────────────

# 조항 본문 최대 관측값 한글 511자 + 출처 프리픽스(~60자) + 여유분
# BAAI/bge-m3: 8192 토큰 한도 → 700자(≈350~500 tok) 는 잘림 없음
MAX_CHUNK_SIZE: int = 700


def _split_at_newlines(chunk: dict, max_size: int) -> list[dict]:
    """청크 텍스트가 max_size를 초과하면 줄 경계에서 분할한다.

    조항 내부 항(①②…) 단위로 자연스럽게 나뉘도록 빈 줄 → 단일 줄 순으로 시도한다.
    분할된 청크는 원본 id에 _p0, _p1, … 접미사를 붙인다.
    """
    text = chunk["text"]
    if len(text) <= max_size:
        return [chunk]

    # 빈 줄 기준으로 먼저 분할, 없으면 단일 줄 기준
    paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
    if len(paragraphs) <= 1:
        paragraphs = [ln.strip() for ln in text.split("\n") if ln.strip()]

    sub_chunks: list[dict] = []
    current: list[str] = []
    current_len = 0

    for para in paragraphs:
        added_len = len(para) + (1 if current else 0)  # 줄바꿈 1자 포함
        if current and current_len + added_len > max_size:
            sub_chunks.append({
                "id": f"{chunk['id']}_p{len(sub_chunks)}",
                "text": "\n".join(current),
                "metadata": chunk["metadata"],
            })
            current, current_len = [para], len(para)
        else:
            current.append(para)
            current_len += added_len

    if current:
        sub_chunks.append({
            "id": f"{chunk['id']}_p{len(sub_chunks)}",
            "text": "\n".join(current),
            "metadata": chunk["metadata"],
        })

    return sub_chunks


# ── 컬렉션 → 정책 문서 제목 매핑 ──
COLLECTION_TO_DOC_TITLE: dict[str, str] = {
    "payment_policy": "주문및결제정책",
    "delivery_policy": "배송정책",
    "return_policy": "반품교환환불정책",
    "quality_policy": "상품품질신선도보증정책",
    "service_policy": "고객서비스운영정책",
    "membership_policy": "개인정보처리및회원정책",
}

# 섹션 헤딩 패턴: "1.", "1.1", "2.3.4" 로 시작하는 줄
_HEADING_RE = re.compile(r"^(\d+(?:\.\d+)*)\s+(.+)$", re.MULTILINE)

# 제X장 / 제X조 패턴 — _apply_heading_markdown 후 PDF/DOCX 공통 형식
#   ## 제N장 제목    (장 — 항상 ## 으로 통일)
#   ### 제N조(제목)  (조 — 항상 ### 으로 통일)
# \s* 대신 [ \t]* 사용 — \s*는 \n까지 소비해 m.start()가 \n 위치를 가리키는 버그 방지
_CHAPTER_RE = re.compile(r"^#{0,6}[ \t]*제(\d+)장[ \t]+(.+)$", re.MULTILINE)
_ARTICLE_RE = re.compile(r"^#{1,6}[ \t]+제(\d+)조\((.+?)\)", re.MULTILINE)

# 머리말 분리용 패턴
# 줄 시작에 있는 단독 장·조 헤딩 — [ \t]*로 \n 소비 방지
_BODY_START_RE = re.compile(r"^#{0,6}[ \t]*제\d+(?:장|조)[\s(]", re.MULTILINE)
# 목차 행 식별: 같은 줄에 제N장/조 + "|" 조합 → TOC 또는 표 헤더
_TOC_INLINE_RE = re.compile(r"제\d+(?:장|조).+\|")


def _split_preamble(text: str) -> tuple[str, str]:
    """파싱된 마크다운에서 머리말(제목·목차·버전)과 본문을 분리한다.

    본문 시작 기준: 줄 단독으로 존재하고 '|' 로 연결되지 않은 첫 번째 제N장/제N조 헤딩.
    목차 행 예: "제1장 배송 안내 | 제2장 배송 지역 | ..." → 건너뜀

    Returns:
        (preamble, body) — preamble은 청킹 대상에서 제외된다.
    """
    for m in _BODY_START_RE.finditer(text):
        pos = m.start()
        line_end = text.find("\n", pos)
        line = text[pos: line_end if line_end != -1 else len(text)]
        if _TOC_INLINE_RE.search(line):
            continue  # 목차 행이므로 건너뜀
        return text[:pos].strip(), text[pos:].strip()

    return "", text  # 본문 시작점 미탐지 시 전체를 본문으로 처리

def chunk_by_sections(
    text: str, source: str, doc_title: str = "", max_size: int = MAX_CHUNK_SIZE
) -> list[dict]:
    """섹션 헤딩 기준으로 텍스트를 청크로 분할.

    doc_title이 주어지면 각 청크 앞에 [doc_title > X.X 섹션명] 출처 프리픽스를 붙여
    LLM이 근거 섹션을 인용할 수 있게 한다.
    max_size 초과 청크는 줄 경계에서 추가 분할된다.

    Returns:
        [{"id": str, "text": str, "metadata": dict}, ...]
    """
    # 헤딩 위치 목록
    matches = list(_HEADING_RE.finditer(text))

    if not matches:
        # 헤딩이 없으면 전체를 하나의 청크로
        prefix = f"[{doc_title}]\n" if doc_title else ""
        raw = {
            "id": f"{source}_chunk_0",
            "text": prefix + text.strip(),
            "metadata": {"source": source, "section": "전체"},
        }
        return _split_at_newlines(raw, max_size)

    chunks = []
    for i, match in enumerate(matches):
        section_num = match.group(1)
        section_title = match.group(2).strip()
        start = match.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)

        content = text[start:end].strip()
        if not content:
            continue

        if doc_title:
            prefix = f"[{doc_title} > {section_num} {section_title}]\n"
            content = prefix + content

        chunk_id = f"{source}_s{section_num.replace('.', '_')}"
        raw = {
            "id": chunk_id,
            "text": content,
            "metadata": {
                "source": source,
                "section": f"{section_num} {section_title}",
                "section_num": section_num,
                **({"doc_title": doc_title} if doc_title else {}),
            },
        }
        chunks.extend(_split_at_newlines(raw, max_size))

    return chunks


def chunk_by_articles(
    text: str, source: str, doc_title: str = "", max_size: int = MAX_CHUNK_SIZE
) -> list[dict]:
    """제X장 > 제X조 체계로 텍스트를 청크로 분할.

    각 청크 앞에 출처 프리픽스를 붙여 LLM이 조·항을 인용할 수 있게 한다.
    max_size 초과 조항은 줄(항) 경계에서 추가 분할된다.

    Returns:
        [{"id": str, "text": str, "metadata": dict}, ...]
    """
    chapters = list(_CHAPTER_RE.finditer(text))
    articles = list(_ARTICLE_RE.finditer(text))

    if not articles:
        return []

    # 장·조 모든 분할점을 위치순으로 정렬
    splits: list[dict] = []
    for m in chapters:
        splits.append({
            "type": "chapter",
            "pos": m.start(),
            "num": m.group(1),
            "title": m.group(2).strip(),
        })
    for m in articles:
        splits.append({
            "type": "article",
            "pos": m.start(),
            "num": m.group(1),
            "title": m.group(2).strip(),
        })
    splits.sort(key=lambda s: s["pos"])

    # 각 조가 속하는 장 계산 + 콘텐츠 범위 추출
    current_chapter = ""
    chunks: list[dict] = []

    for i, sp in enumerate(splits):
        if sp["type"] == "chapter":
            current_chapter = f"제{sp['num']}장 {sp['title']}"
            continue

        # article인 경우
        start = sp["pos"]
        # 다음 분할점(장 또는 조) 시작까지
        end = splits[i + 1]["pos"] if i + 1 < len(splits) else len(text)
        content = text[start:end].strip()
        if not content:
            continue

        # 출처 프리픽스 (doc_title이 없으면 prefix 없이 content만 사용)
        article_label = f"제{sp['num']}조({sp['title']})"
        if doc_title:
            if current_chapter:
                prefix = f"[{doc_title} > {current_chapter} > {article_label}]"
            else:
                prefix = f"[{doc_title} > {article_label}]"
            chunk_text = f"{prefix}\n{content}"
        else:
            chunk_text = content
        chunk_id = f"{source}_art{sp['num']}"

        metadata: dict = {
            "source": source,
            "article": article_label,
            **({"doc_title": doc_title} if doc_title else {}),
        }
        if current_chapter:
            metadata["chapter"] = current_chapter

        raw = {"id": chunk_id, "text": chunk_text, "metadata": metadata}
        chunks.extend(_split_at_newlines(raw, max_size))

    return chunks


# ── ChromaDB 적재 ──────────────────────────────────────────────────────────

def seed_policy_collection(client, ef, filepath: str, collection_name: str) -> int:
    """문서 파일 파싱 → 청킹 → ChromaDB 컬렉션에 적재.

    Returns:
        적재된 청크 수
    """
    filename = os.path.basename(filepath)
    print(f"  파싱 중: {filename}")

    try:
        text = parse_document(filepath)
    except Exception as e:
        print(f"  [오류] 파싱 실패: {e}")
        return 0

    # 파일명에서 소스 식별자 추출 (공백/괄호 제거)
    source = re.sub(r"[\s\(\)]+", "_", os.path.splitext(filename)[0]).strip("_")

    # 머리말(제목·목차·버전) 분리 — 청킹 대상에서 제외
    preamble, body = _split_preamble(text)
    if preamble:
        preview = preamble.replace("\n", " ")[:80]
        print(f"  [머리말 제거] {preview}{'...' if len(preamble) > 80 else ''}")

    # 제X조 패턴이 있으면 조 단위 청킹, 없으면 섹션 청킹 폴백
    # 두 경우 모두 doc_title을 전달하여 출처 프리픽스가 붙게 한다
    doc_title = COLLECTION_TO_DOC_TITLE.get(collection_name, collection_name)
    if _ARTICLE_RE.search(body):
        chunks = chunk_by_articles(body, source, doc_title)
        if not chunks:
            chunks = chunk_by_sections(body, source, doc_title)
    else:
        chunks = chunk_by_sections(body, source, doc_title)

    if not chunks:
        print(f"  [경고] 청크 없음: {filename}")
        return 0

    col = client.get_or_create_collection(name=collection_name, embedding_function=ef)
    col.upsert(
        ids=[c["id"] for c in chunks],
        documents=[c["text"] for c in chunks],
        metadatas=[c["metadata"] for c in chunks],
    )
    count = col.count()
    print(f"  → {collection_name}: {count}개 청크 적재 완료")
    return count


# ── 기존 컬렉션 (JSON 기반) ────────────────────────────────────────────────

def load_json(filename: str) -> list:
    path = os.path.join(DATA_DIR, filename)
    if not os.path.exists(path):
        print(f"[경고] 파일 없음: {path} — 해당 컬렉션을 건너뜁니다.")
        return []
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def seed_faq(client, ef) -> int:
    data = load_json("faq.json")
    if not data:
        return 0

    col = client.get_or_create_collection(name="faq", embedding_function=ef)

    ids, documents, metadatas = [], [], []
    for item in data:
        chunk_text = f"Q: {item['question']}\nA: {item['answer']}"
        ids.append(item["id"])
        documents.append(chunk_text)
        metadatas.append({
            "type": "faq",
            "intent": item["intent"],
            "faq_id": item["id"],
        })

    col.upsert(documents=documents, metadatas=metadatas, ids=ids)
    return col.count()


def seed_storage_guide(client, ef) -> int:
    data = load_json("storage_guide.json")
    if not data:
        return 0

    col = client.get_or_create_collection(name="storage_guide", embedding_function=ef)

    ids, documents, metadatas = [], [], []
    for item in data:
        ids.append(item["id"])
        documents.append(item["guide"])
        metadatas.append({
            "type": "storage",
            "product_name": item["product"],
            "category": item["category"],
        })

    col.upsert(documents=documents, metadatas=metadatas, ids=ids)
    return col.count()


def seed_season_info(client, ef) -> int:
    data = load_json("season_info.json")
    if not data:
        return 0

    col = client.get_or_create_collection(name="season_info", embedding_function=ef)

    ids, documents, metadatas = [], [], []
    for item in data:
        products_text = ", ".join(item["products"]) if item["products"] else "없음"
        doc_text = f"{item['season']} 제철 상품: {products_text}\n{item['description']}"
        ids.append(item["id"])
        documents.append(doc_text)
        metadatas.append({
            "type": "season",
            "season": item["season"],
        })

    col.upsert(documents=documents, metadatas=metadatas, ids=ids)
    return col.count()


def seed_farm_info(client, ef) -> int:
    data = load_json("farm_info.json")
    if not data:
        return 0

    col = client.get_or_create_collection(name="farm_intro", embedding_function=ef)

    ids, documents, metadatas = [], [], []
    for item in data:
        doc_text = f"[{item['category']} > {item['title']}]\n{item['content']}"
        ids.append(item["id"])
        documents.append(doc_text)
        metadatas.append({
            "type": "farm_info",
            "category": item["category"],
            "title": item["title"],
        })

    col.upsert(documents=documents, metadatas=metadatas, ids=ids)
    return col.count()


# ── 진입점 ─────────────────────────────────────────────────────────────────

def find_policy_files() -> list[tuple[str, str]]:
    """DOCS_DIR에서 DOC_TO_COLLECTION 매핑에 해당하는 파일을 찾아 반환.

    Returns:
        [(filepath, collection_name), ...]
    """
    if not os.path.isdir(DOCS_DIR):
        print(f"[경고] 정책 문서 폴더 없음: {DOCS_DIR}")
        return []

    found = []
    for fname in sorted(os.listdir(DOCS_DIR)):
        fpath = os.path.join(DOCS_DIR, fname)
        if not os.path.isfile(fpath):
            continue
        ext = os.path.splitext(fname)[1].lower()
        if ext not in (".pdf", ".docx"):
            continue
        # 파일명 키와 부분 일치 확인
        for key, collection in DOC_TO_COLLECTION.items():
            if key in fname:
                found.append((fpath, collection))
                break
        else:
            print(f"  [건너뜀] 매핑 없음: {fname}")

    return found


def main():
    try:
        import chromadb
    except ImportError:
        print("[오류] chromadb 설치 필요: uv add chromadb")
        sys.exit(1)

    # 기존 chroma_data 전체 삭제 — 개별 컬렉션 삭제 시 HNSW 인덱스 파일이
    # 디스크에 남아 "Nothing found on disk" 오류를 일으키므로 전체 초기화
    import shutil
    if os.path.exists(CHROMA_DIR):
        shutil.rmtree(CHROMA_DIR)
        print(f"기존 ChromaDB 삭제: {CHROMA_DIR}")

    print(f"ChromaDB 초기화 중... ({CHROMA_DIR})")
    client = chromadb.PersistentClient(path=CHROMA_DIR)
    try:
        ef = get_embedding_function()
        print(f"임베딩 provider: {settings.embed_provider}")
    except Exception as e:
        print(f"[오류] 임베딩 함수 초기화 실패: {e}")
        print(f"  → EMBED_PROVIDER={settings.embed_provider} 설정을 확인하세요.")
        sys.exit(1)

    # ── JSON 기반 컬렉션 ──
    print("\n[JSON 컬렉션 적재]")
    faq_count = seed_faq(client, ef)
    print(f"  faq: {faq_count}개")
    storage_count = seed_storage_guide(client, ef)
    print(f"  storage_guide: {storage_count}개")
    season_count = seed_season_info(client, ef)
    print(f"  season_info: {season_count}개")
    farm_count = seed_farm_info(client, ef)
    print(f"  farm_intro: {farm_count}개")

    # ── 정책 문서 적재 ──
    print(f"\n[정책 문서 적재] 경로: {DOCS_DIR}")
    policy_files = find_policy_files()

    if not policy_files:
        print("  적재할 정책 문서 없음.")
    else:
        policy_totals = {}
        for fpath, collection in policy_files:
            count = seed_policy_collection(client, ef, fpath, collection)
            policy_totals[collection] = count

        print("\n정책 컬렉션 요약:")
        for col, cnt in policy_totals.items():
            print(f"  {col}: {cnt}개 청크")

    # ── BM25 인덱스 빌드 ──
    print("\n[BM25 인덱스 빌드]")
    bm25_data = build_bm25_index(client, _ALL_COLLECTIONS)
    with open(BM25_INDEX_PATH, "w", encoding="utf-8") as f:
        json.dump(bm25_data, f, ensure_ascii=False)
    print(f"  저장: {BM25_INDEX_PATH} ({len(bm25_data['ids'])}개 문서)")

    print(f"\n완료. ChromaDB 저장 경로: {os.path.abspath(CHROMA_DIR)}")
    return client, ef


if __name__ == "__main__":
    main()
